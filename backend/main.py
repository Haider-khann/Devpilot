from fastapi import FastAPI, BackgroundTasks, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import sqlite3
import requests
import base64
import ast
import re
import json
import logging
from typing import Dict, List, Optional
from datetime import datetime
from gemini_service import GeminiAIService
from ml_service import CodeMLService

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="DevPilot", version="3.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def init_db():
    conn = sqlite3.connect('devpilot.db')
    cursor = conn.cursor()
    cursor.execute('''CREATE TABLE IF NOT EXISTS repositories (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        github_url TEXT UNIQUE NOT NULL,
        name TEXT, description TEXT, language TEXT,
        stars INTEGER DEFAULT 0, forks INTEGER DEFAULT 0,
        status TEXT DEFAULT 'pending',
        total_files INTEGER DEFAULT 0, total_lines INTEGER DEFAULT 0,
        total_functions INTEGER DEFAULT 0, total_classes INTEGER DEFAULT 0,
        documentation_score REAL DEFAULT 0, code_quality_score REAL DEFAULT 0,
        security_score REAL DEFAULT 0, metrics_json TEXT,
        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
    cursor.execute('''CREATE TABLE IF NOT EXISTS security_issues (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        repository_id INTEGER, severity TEXT, issue_type TEXT,
        description TEXT, file_path TEXT, line_number INTEGER,
        recommendation TEXT, created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)''')
    conn.commit()
    conn.close()

init_db()

class RepositoryCreate(BaseModel):
    github_url: str

class CodeAnalysisRequest(BaseModel):
    code: str
    language: str = "Python"

class GitHubService:
    def __init__(self):
        self.base_url = "https://api.github.com"
        self.headers = {"Accept": "application/vnd.github.v3+json", "User-Agent": "DevPilot"}
    
    def parse_url(self, url):
        url = url.rstrip('/').replace('.git', '')
        parts = url.split('/')
        if 'github.com' not in parts:
            raise ValueError("Invalid GitHub URL")
        idx = parts.index('github.com')
        return parts[idx+1], parts[idx+2]
    
    def get_repo_info(self, url):
        try:
            owner, name = self.parse_url(url)
            resp = requests.get(f"{self.base_url}/repos/{owner}/{name}", headers=self.headers)
            if resp.status_code == 200:
                d = resp.json()
                return {'name': d.get('name'), 'description': d.get('description'), 'language': d.get('language'), 'stars': d.get('stargazers_count', 0), 'forks': d.get('forks_count', 0)}
        except:
            pass
        return {'name': url.split('/')[-1], 'description': None, 'language': None, 'stars': 0, 'forks': 0}
    
    def get_files(self, url, branch='main'):
        try:
            owner, name = self.parse_url(url)
            resp = requests.get(f"{self.base_url}/repos/{owner}/{name}/git/trees/{branch}?recursive=1", headers=self.headers)
            if resp.status_code == 200:
                return [item['path'] for item in resp.json().get('tree', []) if item.get('type') == 'blob']
        except:
            pass
        return []
    
    def get_file_content(self, url, path, branch='main'):
        try:
            owner, name = self.parse_url(url)
            resp = requests.get(f"{self.base_url}/repos/{owner}/{name}/contents/{path}", headers=self.headers)
            if resp.status_code == 200:
                d = resp.json()
                if d.get('encoding') == 'base64':
                    return base64.b64decode(d['content']).decode('utf-8', errors='ignore')
        except:
            pass
        return None

class CodeAnalyzer:
    def analyze_python(self, content):
        try:
            tree = ast.parse(content)
            functions = []
            classes = []
            for node in ast.walk(tree):
                if isinstance(node, ast.FunctionDef):
                    functions.append({'name': node.name, 'docstring': ast.get_docstring(node), 'complexity': self._complexity(node)})
                elif isinstance(node, ast.ClassDef):
                    classes.append({'name': node.name, 'methods': []})
            return {'lines': len(content.splitlines()), 'functions': functions, 'classes': classes, 'num_functions': len(functions), 'num_classes': len(classes)}
        except:
            return self.analyze_generic(content)
    
    def _complexity(self, node):
        c = 1
        for child in ast.walk(node):
            if isinstance(child, (ast.If, ast.While, ast.For, ast.ExceptHandler)):
                c += 1
        return c
    
    def analyze_js(self, content):
        return self.analyze_generic(content)
    
    def analyze_generic(self, content):
        lines = content.splitlines()
        functions = []
        classes = []
        func_patterns = [r'(?:function|def|fn)\s+(\w+)', r'(?:const|let|var)\s+(\w+)\s*=\s*(?:async\s*)?\([^)]*\)\s*=>']
        for pattern in func_patterns:
            for m in re.finditer(pattern, content):
                name = m.group(1)
                if name and name not in [f['name'] for f in functions]:
                    functions.append({'name': name, 'complexity': 1, 'docstring': None})
        class_pattern = r'(?:class|interface|struct|enum)\s+(\w+)'
        for m in re.finditer(class_pattern, content):
            classes.append({'name': m.group(1), 'methods': []})
        return {'lines': len(lines), 'functions': functions, 'classes': classes, 'num_functions': len(functions), 'num_classes': len(classes)}

class SecurityAnalyzer:
    def scan(self, code):
        issues = []
        patterns = [
            ('critical', 'Hardcoded Secret', r'(?i)(password|secret|api[_-]?key|token|auth[_-]?token|client[_-]?secret)\s*=\s*["\'][^"\']+["\']'),
            ('critical', 'AWS Access Key', r'AKIA[0-9A-Z]{16}'),
            ('critical', 'Private Key', r'-----BEGIN (RSA|DSA|EC|OPENSSH) PRIVATE KEY-----'),
            ('high', 'SQL Injection', r'(?i)(execute|cursor\.execute|raw)\s*\(.*f["\']'),
            ('high', 'Command Injection', r'(?i)(os\.system|subprocess\.(call|run|Popen)|eval|exec|shell=True)\('),
            ('high', 'XSS Vulnerability', r'(?i)(innerHTML|document\.write|dangerouslySetInnerHTML)'),
            ('medium', 'Debug Mode', r'(?i)debug\s*=\s*True'),
            ('medium', 'Insecure Deserialization', r'(?i)(pickle\.loads|yaml\.load\()'),
            ('medium', 'Insecure HTTP', r'http://(?!localhost|127\.0\.0\.1)'),
            ('medium', 'Weak Crypto', r'(?i)(md5|sha1)\('),
            ('low', 'Bare Except', r'except\s*:'),
            ('low', 'TODO Comment', r'#\s*TODO'),
            ('low', 'Print Statement', r'^\s*print\('),
            ('low', 'Console Log', r'console\.log\('),
        ]
        for line_num, line in enumerate(code.splitlines(), 1):
            for severity, issue_type, pattern in patterns:
                if re.search(pattern, line):
                    issues.append({'severity': severity, 'type': issue_type, 'line': line_num, 'description': line.strip()[:100], 'recommendation': self._rec(issue_type)})
        return issues
    
    def _rec(self, t):
        recs = {
            'Hardcoded Secret': 'Use environment variables', 'AWS Access Key': 'Rotate key, use IAM roles',
            'Private Key': 'Never commit private keys', 'SQL Injection': 'Use parameterized queries',
            'Command Injection': 'Avoid shell=True', 'XSS Vulnerability': 'Use textContent',
            'Debug Mode': 'Disable in production', 'Insecure Deserialization': 'Use safe_load',
            'Insecure HTTP': 'Use HTTPS', 'Weak Crypto': 'Use SHA-256 or bcrypt',
            'Bare Except': 'Catch specific exceptions', 'TODO Comment': 'Complete or remove',
            'Print Statement': 'Use logging', 'Console Log': 'Remove or use logger',
        }
        return recs.get(t, 'Review and fix')

def analyze_repository(repo_id, github_url):
    conn = sqlite3.connect('devpilot.db')
    cursor = conn.cursor()
    try:
        cursor.execute("UPDATE repositories SET status='analyzing' WHERE id=?", (repo_id,))
        conn.commit()
        github = GitHubService()
        analyzer = CodeAnalyzer()
        security = SecurityAnalyzer()
        repo_info = github.get_repo_info(github_url)
        files = github.get_files(github_url)
        all_exts = ('.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.go', '.rs', '.cpp', '.c', '.h', '.rb', '.php', '.html', '.css', '.json', '.xml', '.yaml', '.yml', '.md', '.sql', '.sh', '.toml', '.ini', '.cfg', '.txt')
        code_files = [f for f in files if f.endswith(all_exts)][:200]
        total_lines = 0
        total_functions = 0
        total_classes = 0
        all_functions = []
        all_security_issues = []
        for file_path in code_files:
            content = github.get_file_content(github_url, file_path)
            if content:
                if file_path.endswith('.py'):
                    result = analyzer.analyze_python(content)
                elif file_path.endswith(('.js', '.ts', '.jsx', '.tsx')):
                    result = analyzer.analyze_js(content)
                else:
                    result = analyzer.analyze_generic(content)
                total_lines += result['lines']
                total_functions += result['num_functions']
                total_classes += result['num_classes']
                all_functions.extend(result['functions'])
                for issue in security.scan(content):
                    issue['file_path'] = file_path
                    all_security_issues.append(issue)
        # Better scoring based on actual content analysis
        total_comments = 0
        total_code_lines = 0
        for file_path in code_files:
            content = github.get_file_content(github_url, file_path)
            if content:
                lines = content.splitlines()
                for line in lines:
                    stripped = line.strip()
                    if stripped.startswith('#') or stripped.startswith('//') or stripped.startswith('<!--') or stripped.startswith('/*') or stripped.startswith('*'):
                        total_comments += 1
                    elif stripped:
                        total_code_lines += 1
        
        # Documentation score: based on comments ratio AND function docstrings
        documented = sum(1 for f in all_functions if f.get('docstring'))
        func_doc_score = (documented / len(all_functions) * 100) if all_functions else 0
        comment_ratio = (total_comments / total_code_lines * 100) if total_code_lines > 0 else 0
        doc_score = round((func_doc_score * 0.6 + comment_ratio * 0.4), 2)
        if doc_score == 0 and not all_functions:
            doc_score = round(min(comment_ratio, 100), 2)
        
        # Quality score: based on complexity, file size, and structure
        complex_funcs = sum(1 for f in all_functions if f.get('complexity', 1) > 10)
        complexity_penalty = (complex_funcs / len(all_functions) * 40) if all_functions else 0
        avg_lines_per_file = total_lines / len(code_files) if code_files else 0
        size_penalty = min(20, max(0, (avg_lines_per_file - 300) / 20))
        structure_bonus = min(20, total_classes * 2 + total_functions * 0.5)
        quality_score = round(max(30, min(100, 80 - complexity_penalty - size_penalty + structure_bonus)), 2)
        
        # Security score: based on issues AND file types
        critical_issues = sum(1 for i in all_security_issues if i['severity'] == 'critical')
        high_issues = sum(1 for i in all_security_issues if i['severity'] == 'high')
        medium_issues = sum(1 for i in all_security_issues if i['severity'] == 'medium')
        low_issues = sum(1 for i in all_security_issues if i['severity'] == 'low')
        issue_penalty = critical_issues * 20 + high_issues * 10 + medium_issues * 5 + low_issues * 2
        security_score = round(max(0, 100 - issue_penalty), 2)
        metrics = {'total_files': len(code_files), 'total_lines': total_lines, 'total_functions': total_functions, 'total_classes': total_classes, 'documentation_score': round(doc_score, 2), 'code_quality_score': round(quality_score, 2), 'security_score': round(security_score, 2), 'security_issues_count': len(all_security_issues)}
        cursor.execute("""UPDATE repositories SET name=?, description=?, language=?, stars=?, forks=?, status='completed', total_files=?, total_lines=?, total_functions=?, total_classes=?, documentation_score=?, code_quality_score=?, security_score=?, metrics_json=?, updated_at=CURRENT_TIMESTAMP WHERE id=?""", (repo_info['name'], repo_info['description'], repo_info['language'], repo_info['stars'], repo_info['forks'], metrics['total_files'], metrics['total_lines'], metrics['total_functions'], metrics['total_classes'], metrics['documentation_score'], metrics['code_quality_score'], metrics['security_score'], json.dumps(metrics), repo_id))
        conn.commit()
        for issue in all_security_issues:
            cursor.execute("INSERT INTO security_issues (repository_id, severity, issue_type, description, file_path, line_number, recommendation) VALUES (?, ?, ?, ?, ?, ?, ?)", (repo_id, issue['severity'], issue['type'], issue['description'], issue.get('file_path', ''), issue['line'], issue['recommendation']))
        conn.commit()
        logger.info(f"Analysis done: {metrics}")
    except Exception as e:
        logger.error(f"Failed: {e}")
        cursor.execute("UPDATE repositories SET status='failed' WHERE id=?", (repo_id,))
        conn.commit()
    finally:
        conn.close()

@app.get("/")
async def root():
    return {"name": "DevPilot", "version": "3.0.0"}

@app.get("/health")
async def health():
    return {"status": "healthy"}

@app.get("/api/repositories")
async def list_repositories():
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM repositories ORDER BY created_at DESC")
    repos = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return repos

@app.post("/api/repositories")
async def add_repository(repo: RepositoryCreate, background_tasks: BackgroundTasks):
    conn = sqlite3.connect('devpilot.db')
    cursor = conn.cursor()
    name = repo.github_url.rstrip('/').split('/')[-1]
    try:
        cursor.execute("INSERT INTO repositories (github_url, name) VALUES (?, ?)", (repo.github_url, name))
        conn.commit()
        repo_id = cursor.lastrowid
        background_tasks.add_task(analyze_repository, repo_id, repo.github_url)
        return {"id": repo_id, "status": "pending"}
    except sqlite3.IntegrityError:
        raise HTTPException(status_code=400, detail="Repository already exists")
    finally:
        conn.close()

@app.get("/api/repositories/{repo_id}")
async def get_repository(repo_id: int):
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM repositories WHERE id=?", (repo_id,))
    row = cursor.fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Not found")
    return dict(row)

@app.delete("/api/repositories/{repo_id}")
async def delete_repository(repo_id: int):
    conn = sqlite3.connect('devpilot.db')
    cursor = conn.cursor()
    cursor.execute("DELETE FROM repositories WHERE id=?", (repo_id,))
    cursor.execute("DELETE FROM security_issues WHERE repository_id=?", (repo_id,))
    conn.commit()
    conn.close()
    return {"message": "Deleted"}

@app.get("/api/repositories/{repo_id}/security")
async def get_security_issues(repo_id: int):
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM security_issues WHERE repository_id=? ORDER BY severity DESC", (repo_id,))
    issues = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return issues

@app.post("/api/analyze-code")
async def analyze_code(request: CodeAnalysisRequest):
    security = SecurityAnalyzer()
    issues = security.scan(request.code)
    severity_counts = {'critical': 0, 'high': 0, 'medium': 0, 'low': 0}
    for issue in issues:
        severity_counts[issue['severity']] += 1
    return {"security_issues": issues, "severity_counts": severity_counts, "has_issues": len(issues) > 0, "code_length": len(request.code.splitlines())}

@app.post("/api/compare")
async def compare_repositories(request: dict):
    url1 = request.get('repo1_url', '')
    url2 = request.get('repo2_url', '')
    if not url1 or not url2:
        raise HTTPException(status_code=400, detail="Both repo URLs required")
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM repositories WHERE github_url=?", (url1,))
    repo1 = cursor.fetchone()
    cursor.execute("SELECT * FROM repositories WHERE github_url=?", (url2,))
    repo2 = cursor.fetchone()
    conn.close()
    if not repo1 or not repo2:
        raise HTTPException(status_code=404, detail="One or both repos not found")
    repo1 = dict(repo1)
    repo2 = dict(repo2)
    fields = ['total_files', 'total_lines', 'total_functions', 'total_classes', 'documentation_score', 'code_quality_score', 'security_score']
    comparison = {}
    for field in fields:
        v1 = repo1.get(field) or 0
        v2 = repo2.get(field) or 0
        comparison[field] = {'repo1': v1, 'repo2': v2, 'difference': v2 - v1, 'winner': 'repo1' if v1 > v2 else 'repo2' if v2 > v1 else 'tie'}
    return {'repo1': {'name': repo1['name'], 'url': repo1['github_url']}, 'repo2': {'name': repo2['name'], 'url': repo2['github_url']}, 'comparison': comparison, 'overall_winner': 'repo1' if sum(1 for f in fields if (repo1.get(f) or 0) > (repo2.get(f) or 0)) > len(fields)//2 else 'repo2'}

@app.get("/api/export/{repo_id}/json")
async def export_json(repo_id: int):
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM repositories WHERE id=?", (repo_id,))
    repo = cursor.fetchone()
    if not repo:
        conn.close()
        raise HTTPException(status_code=404, detail="Not found")
    repo = dict(repo)
    cursor.execute("SELECT * FROM security_issues WHERE repository_id=?", (repo_id,))
    issues = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {'repository': repo, 'security_issues': issues, 'exported_at': datetime.now().isoformat()}

@app.get("/api/stats")
async def get_stats():
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) as c FROM repositories")
    total_repos = cursor.fetchone()['c']
    cursor.execute("SELECT COUNT(*) as c FROM repositories WHERE status='completed'")
    completed = cursor.fetchone()['c']
    cursor.execute("SELECT COUNT(*) as c FROM repositories WHERE status='analyzing'")
    analyzing = cursor.fetchone()['c']
    cursor.execute("SELECT AVG(documentation_score) as a FROM repositories WHERE status='completed'")
    avg_doc = cursor.fetchone()['a'] or 0
    cursor.execute("SELECT AVG(code_quality_score) as a FROM repositories WHERE status='completed'")
    avg_quality = cursor.fetchone()['a'] or 0
    cursor.execute("SELECT AVG(security_score) as a FROM repositories WHERE status='completed'")
    avg_security = cursor.fetchone()['a'] or 0
    cursor.execute("SELECT SUM(total_lines) as s FROM repositories")
    sum_lines = cursor.fetchone()['s'] or 0
    cursor.execute("SELECT SUM(total_functions) as s FROM repositories")
    sum_funcs = cursor.fetchone()['s'] or 0
    conn.close()
    return {'total_repositories': total_repos, 'completed': completed, 'analyzing': analyzing, 'average_documentation_score': round(avg_doc, 2), 'average_quality_score': round(avg_quality, 2), 'average_security_score': round(avg_security, 2), 'total_lines_analyzed': sum_lines, 'total_functions_found': sum_funcs}


@app.get("/api/repositories/{repo_id}/files")
async def get_repo_files(repo_id: int):
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM repositories WHERE id=?", (repo_id,))
    repo = cursor.fetchone()
    if not repo:
        conn.close()
        raise HTTPException(status_code=404, detail="Not found")
    conn.close()
    
    # Fetch files from GitHub
    github = GitHubService()
    files = github.get_files(dict(repo)['github_url'])
    all_exts = ('.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.go', '.rs', '.cpp', '.c', '.h', '.rb', '.php', '.html', '.css', '.json', '.xml', '.yaml', '.yml', '.md', '.sql', '.sh', '.toml', '.ini', '.cfg', '.txt')
    code_files = [f for f in files if f.endswith(all_exts)][:200]
    
    return {"repository_id": repo_id, "files": code_files, "total": len(code_files)}

@app.post("/api/repositories/{repo_id}/reanalyze")
async def reanalyze_repository(repo_id: int, background_tasks: BackgroundTasks):
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM repositories WHERE id=?", (repo_id,))
    repo = cursor.fetchone()
    conn.close()
    
    if not repo:
        raise HTTPException(status_code=404, detail="Not found")
    
    repo = dict(repo)
    background_tasks.add_task(analyze_repository, repo_id, repo['github_url'])
    return {"message": "Re-analysis started", "status": "analyzing"}


@app.post("/api/ai/summarize")
async def ai_summarize(request: dict):
    code = request.get('code', '')
    language = request.get('language', 'Python')
    
    if not code:
        raise HTTPException(status_code=400, detail="Code required")
    
    lines = code.splitlines()
    total_lines = len(lines)
    
    # Pattern-based analysis
    functions = re.findall(r'(?:def|function)\s+(\w+)', code)
    classes = re.findall(r'class\s+(\w+)', code)
    imports = re.findall(r'(?:import|from)\s+([\w.]+)', code)
    comments = [l for l in lines if l.strip().startswith('#') or l.strip().startswith('//')]
    
    # Generate summary
    summary_parts = []
    summary_parts.append(f"This {language} code contains {total_lines} lines.")
    
    if functions:
        summary_parts.append(f"It defines {len(functions)} functions: {', '.join(functions[:5])}.")
    if classes:
        summary_parts.append(f"It defines {len(classes)} classes: {', '.join(classes[:5])}.")
    if imports:
        summary_parts.append(f"It imports {len(imports)} modules: {', '.join(imports[:5])}.")
    if comments:
        summary_parts.append(f"It has {len(comments)} comments ({round(len(comments)/total_lines*100, 1)}% comment ratio).")
    
    # Code complexity estimate
    control_flow = len(re.findall(r'(if|for|while|switch|case)', code))
    if control_flow > 10:
        summary_parts.append(f"It has {control_flow} control flow statements (high complexity).")
    elif control_flow > 0:
        summary_parts.append(f"It has {control_flow} control flow statements (moderate complexity).")
    else:
        summary_parts.append("It has low complexity.")
    
    return {
        "summary": " ".join(summary_parts),
        "metrics": {
            "lines": total_lines,
            "functions": len(functions),
            "classes": len(classes),
            "imports": len(imports),
            "comments": len(comments),
            "control_flow": control_flow
        }
    }

@app.post("/api/ai/improve")
async def ai_improve(request: dict):
    code = request.get('code', '')
    if not code:
        raise HTTPException(status_code=400, detail="Code required")
    
    suggestions = []
    
    # Check for common improvements
    if 'print(' in code:
        suggestions.append("Replace print statements with logging for production code")
    if 'except:' in code and 'except Exception' not in code:
        suggestions.append("Catch specific exceptions instead of bare except")
    if 'TODO' in code:
        suggestions.append("Complete or remove TODO comments")
    if 'pass' in code:
        suggestions.append("Implement placeholder 'pass' statements")
    if 'global ' in code:
        suggestions.append("Avoid using global variables")
    if 'eval(' in code:
        suggestions.append("Avoid using eval() - it's a security risk")
    if '=' * 10 in code:
        suggestions.append("Break long lines into multiple lines")
    
    if not suggestions:
        suggestions.append("Code looks good! No major improvements needed.")
    
    return {"suggestions": suggestions, "count": len(suggestions)}

@app.post("/api/ai/generate-tests")
async def ai_generate_tests(request: dict):
    code = request.get('code', '')
    function_name = request.get('function_name', 'my_function')
    if not code:
        raise HTTPException(status_code=400, detail="Code required")
    
    # Auto-detect function name from code
    detected = re.findall(r'(?:def|function)\s+(\w+)', code)
    if detected:
        function_name = detected[0]
    
    # Generate basic test template
    test_code = f'''import pytest

def test_{function_name}_normal():
    """Test normal case."""
    # Arrange
    # Act
    # Assert
    pass

def test_{function_name}_edge_case():
    """Test edge case."""
    pass

def test_{function_name}_error():
    """Test error handling."""
    with pytest.raises(Exception):
        pass
'''
    
    return {
        "test_code": test_code,
        "test_framework": "pytest",
        "test_count": 3
    }


# AI Chat Endpoints
@app.post("/api/ai/chat")
async def ai_chat(request: dict):
    message = request.get('message', '')
    code_context = request.get('code', '')
    
    if not message:
        raise HTTPException(status_code=400, detail="Message required")
    
    ai = GeminiAIService()
    if code_context:
        response = ai.answer_question(code_context, message)
    else:
        response = ai.chat("You are DevPilot AI assistant. Help with coding questions.", message)
    
    return {"response": response}

@app.post("/api/ai/improve-code")
async def ai_improve_code(request: dict):
    code = request.get('code', '')
    instruction = request.get('instruction', '')
    
    if not code:
        raise HTTPException(status_code=400, detail="Code required")
    
    ai = GeminiAIService()
    improved = ai.improve_code(code, instruction)
    
    return {"improved_code": improved}

@app.post("/api/ai/generate-code")
async def ai_generate_code(request: dict):
    description = request.get('description', '')
    reference_code = request.get('reference_code', '')
    
    if not description:
        raise HTTPException(status_code=400, detail="Description required")
    
    ai = GeminiAIService()
    generated = ai.generate_code(description, reference_code)
    
    return {"generated_code": generated}

@app.post("/api/ai/summarize-code")
async def ai_summarize_code(request: dict):
    code = request.get('code', '')
    language = request.get('language', 'Python')
    
    if not code:
        raise HTTPException(status_code=400, detail="Code required")
    
    ai = GeminiAIService()
    summary = ai.summarize_code(code, language)
    
    return {"summary": summary}


@app.get("/api/reports/analysis/{repo_id}")
async def get_analysis_report(repo_id: int):
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM repositories WHERE id=?", (repo_id,))
    repo = cursor.fetchone()
    if not repo:
        conn.close()
        raise HTTPException(status_code=404, detail="Not found")
    repo = dict(repo)
    cursor.execute("SELECT * FROM security_issues WHERE repository_id=?", (repo_id,))
    issues = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {
        "report_type": "analysis",
        "repository": repo,
        "security_issues": issues,
        "generated_at": datetime.now().isoformat()
    }

@app.get("/api/reports/statistics")
async def get_statistics_report():
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) as c FROM repositories")
    total = cursor.fetchone()['c']
    cursor.execute("SELECT COUNT(*) as c FROM repositories WHERE status='completed'")
    completed = cursor.fetchone()['c']
    cursor.execute("SELECT AVG(documentation_score) as a FROM repositories")
    avg_doc = cursor.fetchone()['a'] or 0
    cursor.execute("SELECT AVG(code_quality_score) as a FROM repositories")
    avg_quality = cursor.fetchone()['a'] or 0
    cursor.execute("SELECT AVG(security_score) as a FROM repositories")
    avg_security = cursor.fetchone()['a'] or 0
    cursor.execute("SELECT SUM(total_lines) as s FROM repositories")
    total_lines = cursor.fetchone()['s'] or 0
    cursor.execute("SELECT SUM(total_files) as s FROM repositories")
    total_files = cursor.fetchone()['s'] or 0
    cursor.execute("SELECT language, COUNT(*) as c FROM repositories GROUP BY language")
    languages = [dict(row) for row in cursor.fetchall()]
    conn.close()
    return {
        "report_type": "statistics",
        "total_repositories": total,
        "completed": completed,
        "average_documentation": round(avg_doc, 2),
        "average_quality": round(avg_quality, 2),
        "average_security": round(avg_security, 2),
        "total_lines_analyzed": total_lines,
        "total_files_analyzed": total_files,
        "languages": languages,
        "generated_at": datetime.now().isoformat()
    }

@app.get("/api/reports/comparison/{repo1_id}/{repo2_id}")
async def get_comparison_report(repo1_id: int, repo2_id: int):
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM repositories WHERE id=?", (repo1_id,))
    r1 = cursor.fetchone()
    cursor.execute("SELECT * FROM repositories WHERE id=?", (repo2_id,))
    r2 = cursor.fetchone()
    conn.close()
    if not r1 or not r2:
        raise HTTPException(status_code=404, detail="Not found")
    r1, r2 = dict(r1), dict(r2)
    fields = ['total_files','total_lines','total_functions','total_classes','documentation_score','code_quality_score','security_score']
    comparison = {}
    for f in fields:
        v1, v2 = r1.get(f) or 0, r2.get(f) or 0
        comparison[f] = {'repo1': v1, 'repo2': v2, 'winner': 'repo1' if v1 > v2 else 'repo2' if v2 > v1 else 'tie'}
    return {
        "report_type": "comparison",
        "repo1": {"name": r1['name'], "url": r1['github_url']},
        "repo2": {"name": r2['name'], "url": r2['github_url']},
        "comparison": comparison,
        "generated_at": datetime.now().isoformat()
    }


@app.get("/api/reports/analysis/{repo_id}/pdf")
async def download_analysis_pdf(repo_id: int):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    import io
    
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM repositories WHERE id=?", (repo_id,))
    repo = cursor.fetchone()
    if not repo:
        conn.close()
        raise HTTPException(status_code=404, detail="Not found")
    repo = dict(repo)
    cursor.execute("SELECT * FROM security_issues WHERE repository_id=?", (repo_id,))
    issues = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('Title', parent=styles['Heading1'], fontSize=24, textColor=colors.HexColor('#6366f1'), spaceAfter=20)
    heading_style = ParagraphStyle('Heading', parent=styles['Heading2'], fontSize=16, textColor=colors.HexColor('#4f46e5'), spaceBefore=15, spaceAfter=10)
    body_style = ParagraphStyle('Body', parent=styles['Normal'], fontSize=11, spaceAfter=6)
    
    story = []
    story.append(Paragraph("DevPilot - Repository Analysis Report", title_style))
    story.append(Spacer(1, 10))
    story.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", body_style))
    story.append(Spacer(1, 20))
    
    story.append(Paragraph("Repository Information", heading_style))
    info_data = [
        ["Name", repo.get('name', 'N/A')],
        ["GitHub URL", repo.get('github_url', 'N/A')],
        ["Description", repo.get('description', 'N/A') or 'N/A'],
        ["Language", repo.get('language', 'N/A') or 'N/A'],
        ["Stars", str(repo.get('stars', 0))],
        ["Status", repo.get('status', 'N/A')],
    ]
    info_table = Table(info_data, colWidths=[100, 350])
    info_table.setStyle(TableStyle([('BACKGROUND', (0,0), (0,-1), colors.HexColor('#f1f5f9')), ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')), ('PADDING', (0,0), (-1,-1), 8)]))
    story.append(info_table)
    story.append(Spacer(1, 15))
    
    story.append(Paragraph("Code Metrics", heading_style))
    metrics_data = [
        ["Metric", "Value"],
        ["Total Files", str(repo.get('total_files', 0))],
        ["Total Lines", str(repo.get('total_lines', 0))],
        ["Total Functions", str(repo.get('total_functions', 0))],
        ["Total Classes", str(repo.get('total_classes', 0))],
        ["Documentation Score", f"{repo.get('documentation_score', 0)}%"],
        ["Code Quality Score", f"{repo.get('code_quality_score', 0)}%"],
        ["Security Score", f"{repo.get('security_score', 0)}%"],
    ]
    metrics_table = Table(metrics_data, colWidths=[200, 250])
    metrics_table.setStyle(TableStyle([('BACKGROUND', (0,0), (-1,0), colors.HexColor('#6366f1')), ('TEXTCOLOR', (0,0), (-1,0), colors.white), ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')), ('PADDING', (0,0), (-1,-1), 8)]))
    story.append(metrics_table)
    story.append(Spacer(1, 15))
    
    story.append(Paragraph(f"Security Issues ({len(issues)} found)", heading_style))
    if issues:
        for issue in issues:
            story.append(Paragraph(f"<b>{issue['issue_type']}</b> ({issue['severity']}) - {issue.get('file_path', 'N/A')} line {issue.get('line_number', 'N/A')}", body_style))
            story.append(Paragraph(f"Recommendation: {issue.get('recommendation', 'N/A')}", body_style))
            story.append(Spacer(1, 5))
    else:
        story.append(Paragraph("No security issues found.", body_style))
    
    doc.build(story)
    buffer.seek(0)
    
    from fastapi.responses import StreamingResponse
    return StreamingResponse(buffer, media_type="application/pdf", headers={"Content-Disposition": f"attachment; filename=devpilot-report-{repo_id}.pdf"})

@app.get("/api/reports/analysis/{repo_id}/word")
async def download_analysis_word(repo_id: int):
    from docx import Document
    import io
    
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM repositories WHERE id=?", (repo_id,))
    repo = cursor.fetchone()
    if not repo:
        conn.close()
        raise HTTPException(status_code=404, detail="Not found")
    repo = dict(repo)
    cursor.execute("SELECT * FROM security_issues WHERE repository_id=?", (repo_id,))
    issues = [dict(row) for row in cursor.fetchall()]
    conn.close()
    
    doc = Document()
    doc.add_heading('DevPilot - Repository Analysis Report', 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    
    doc.add_heading('Repository Information', level=1)
    doc.add_paragraph(f"Name: {repo.get('name', 'N/A')}")
    doc.add_paragraph(f"GitHub URL: {repo.get('github_url', 'N/A')}")
    doc.add_paragraph(f"Description: {repo.get('description', 'N/A') or 'N/A'}")
    doc.add_paragraph(f"Language: {repo.get('language', 'N/A') or 'N/A'}")
    doc.add_paragraph(f"Stars: {repo.get('stars', 0)}")
    doc.add_paragraph(f"Status: {repo.get('status', 'N/A')}")
    
    doc.add_heading('Code Metrics', level=1)
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Light Grid Accent 1'
    metrics = [
        ('Metric', 'Value'),
        ('Total Files', str(repo.get('total_files', 0))),
        ('Total Lines', str(repo.get('total_lines', 0))),
        ('Total Functions', str(repo.get('total_functions', 0))),
        ('Total Classes', str(repo.get('total_classes', 0))),
        ('Documentation Score', f"{repo.get('documentation_score', 0)}%"),
        ('Code Quality Score', f"{repo.get('code_quality_score', 0)}%"),
        ('Security Score', f"{repo.get('security_score', 0)}%"),
    ]
    for i, (k, v) in enumerate(metrics):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_heading(f'Security Issues ({len(issues)} found)', level=1)
    if issues:
        for issue in issues:
            doc.add_paragraph(f"{issue['issue_type']} ({issue['severity']})", style='Intense Quote')
            doc.add_paragraph(f"File: {issue.get('file_path', 'N/A')} Line: {issue.get('line_number', 'N/A')}")
            doc.add_paragraph(f"Recommendation: {issue.get('recommendation', 'N/A')}")
    else:
        doc.add_paragraph("No security issues found.")
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    from fastapi.responses import StreamingResponse
    return StreamingResponse(buffer, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename=devpilot-report-{repo_id}.docx"})


@app.post("/api/ai/chat-smart")
async def ai_chat_smart(request: dict):
    message = request.get('message', '')
    mode = request.get('mode', 'general')  # 'support' or 'general'
    code_context = request.get('code', '')
    
    if not message:
        raise HTTPException(status_code=400, detail="Message required")
    
    ai = GeminiAIService()
    
    if mode == 'support':
        system_prompt = """You are DevPilot Support Assistant. Help users with:
1. How to use DevPilot platform
2. Analyzing repositories
3. Security scanning
4. Comparing repos
5. Reports and downloads
6. AI tools usage
7. Troubleshooting issues

DevPilot Features:
- Repository Analysis: Paste GitHub URL to analyze code
- Security Scanner: Detects 15 vulnerability patterns
- Compare: Side-by-side repo comparison
- Reports: PDF, Word, JSON downloads
- AI Tools: Summarize, Improve, Generate code
- Chat Bot: Ask coding questions

Contact: support@devpilot.io
"""
    else:
        system_prompt = """You are DevPilot AI, a powerful coding assistant. You can:
1. Answer programming questions
2. Write code from scratch
3. Debug and fix code
4. Explain complex concepts
5. Suggest improvements
6. Generate full projects
7. Help with any language/framework

Be thorough and provide complete, working code when asked."""
    
    if code_context:
        prompt = f"{system_prompt}\n\nContext code:\n```\n{code_context[:3000]}\n```\n\nUser: {message}"
    else:
        prompt = f"{system_prompt}\n\nUser: {message}"
    
    response = ai.chat(system_prompt, prompt, max_tokens=2000)
    return {"response": response, "mode": mode}


# ============ ML ENDPOINTS ============
@app.post("/api/ml/train")
async def train_model():
    ml = CodeMLService()
    code_samples, languages = ml.get_sample_training_data()
    result = ml.train_language_classifier(code_samples, languages)
    return {"status": "success", "message": "Model trained successfully", **result}

@app.post("/api/ml/predict-language")
async def predict_language(request: dict):
    code = request.get('code', '')
    if not code:
        raise HTTPException(status_code=400, detail="Code required")
    ml = CodeMLService()
    result = ml.predict_language(code)
    return result

@app.get("/api/ml/status")
async def ml_status():
    import os
    model_file = os.path.join('ml_models', 'language_classifier.pkl')
    return {"trained": os.path.exists(model_file), "model_path": model_file}


@app.get("/api/repositories/{repo_id}/ml-analysis")
async def get_ml_analysis(repo_id: int):
    conn = sqlite3.connect('devpilot.db')
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM repositories WHERE id=?", (repo_id,))
    repo = cursor.fetchone()
    conn.close()
    
    if not repo:
        raise HTTPException(status_code=404, detail="Not found")
    
    repo = dict(repo)
    github = GitHubService()
    ml = CodeMLService()
    
    files = github.get_files(repo['github_url'])
    code_files = [f for f in files if f.endswith(('.py','.js','.ts','.html','.css','.json','.md','.txt','.xml','.yaml','.yml'))][:100]
    
    file_analysis = []
    language_stats = {}
    
    for file_path in code_files:
        content = github.get_file_content(repo['github_url'], file_path)
        if content:
            # ML prediction
            prediction = ml.predict_language(content)
            predicted_lang = prediction.get('predicted_language', 'unknown')
            confidence = prediction.get('confidence', 0)
            
            # Count lines
            lines = len(content.splitlines())
            
            # Quality score (simple heuristic + ML confidence)
            quality = 100 if confidence > 80 else (60 if confidence > 50 else 30)
            
            # Track language stats
            if predicted_lang not in language_stats:
                language_stats[predicted_lang] = {'files': 0, 'lines': 0}
            language_stats[predicted_lang]['files'] += 1
            language_stats[predicted_lang]['lines'] += lines
            
            file_analysis.append({
                'file_path': file_path,
                'predicted_language': predicted_lang,
                'confidence': confidence,
                'lines': lines,
                'quality_score': quality,
            })
    
    return {
        'repository': repo['name'],
        'total_files_analyzed': len(file_analysis),
        'language_stats': language_stats,
        'files': file_analysis
    }


@app.post("/api/ml/train-quality")
async def train_quality_model():
    ml = CodeMLService()
    result = ml.train_quality_model()
    return {"status": "success", **result}

@app.post("/api/ml/predict-quality")
async def predict_quality(request: dict):
    code = request.get('code', '')
    if not code:
        raise HTTPException(status_code=400, detail="Code required")
    ml = CodeMLService()
    result = ml.predict_quality(code)
    return result


@app.post("/api/ml/similarity")
async def check_similarity(request: dict):
    code1 = request.get('code1', '')
    code2 = request.get('code2', '')
    if not code1 or not code2:
        raise HTTPException(status_code=400, detail="Both code samples required")
    ml = CodeMLService()
    similarity = ml.find_similar_code(code1, code2)
    is_dup = bool(similarity > 60)
    return {"similarity": similarity, "is_duplicate": is_dup}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)